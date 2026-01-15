"""
Transcription Result Card Widget - Display individual transcription result
"""

import customtkinter as ctk
from tkinter import messagebox, filedialog
import os


class TranscriptionResultCard(ctk.CTkFrame):
    """Widget for displaying a single transcription result"""
    
    def __init__(self, parent, filename, text, **kwargs):
        super().__init__(parent, corner_radius=10, **kwargs)
        
        self.filename = filename
        self.text = text
        
        # Configure grid
        self.grid_columnconfigure(0, weight=1)
        
        # Header with filename and buttons
        header_frame = ctk.CTkFrame(self, fg_color="transparent")
        header_frame.grid(row=0, column=0, sticky="ew", padx=15, pady=(15, 5))
        header_frame.grid_columnconfigure(0, weight=1)
        
        # Filename label
        filename_label = ctk.CTkLabel(
            header_frame,
            text=f"📄 {filename}",
            font=("Arial", 13, "bold"),
            anchor="w"
        )
        filename_label.grid(row=0, column=0, sticky="w")
        
        # Export button
        export_btn = ctk.CTkButton(
            header_frame,
            text="Export",
            width=80,
            height=28,
            corner_radius=6,
            command=self.export_text
        )
        export_btn.grid(row=0, column=1, padx=(0, 10))
        
        # Copy button
        copy_btn = ctk.CTkButton(
            header_frame,
            text="Copy",
            width=70,
            height=28,
            corner_radius=6,
            command=self.copy_text
        )
        copy_btn.grid(row=0, column=2)
        
        # Text display (read-only)
        self.text_box = ctk.CTkTextbox(
            self,
            font=("Arial", 11),
            corner_radius=8,
            wrap="word",
            height=150
        )
        self.text_box.grid(row=1, column=0, sticky="nsew", padx=15, pady=(0, 15))
        self.text_box.insert("1.0", text)
        self.text_box.configure(state="disabled")
        
        # Configure row weight for text box
        self.grid_rowconfigure(1, weight=1)
    
    def copy_text(self):
        """Copy text to clipboard"""
        try:
            self.clipboard_clear()
            self.clipboard_append(self.text)
            messagebox.showinfo("Copied", "Text copied to clipboard!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to copy text: {str(e)}")
    
    def export_text(self):
        """Export text to .txt or .docx file"""
        # Ask for file type
        file_types = [
            ("Text files", "*.txt"),
            ("Word documents", "*.docx")
        ]
        
        # Suggest filename based on original
        base_name = os.path.splitext(self.filename)[0]
        default_name = f"{base_name}_transcription.txt"
        
        filepath = filedialog.asksaveasfilename(
            title="Export Transcription",
            defaultextension=".txt",
            filetypes=file_types,
            initialfile=default_name
        )
        
        if not filepath:
            return
        
        try:
            # Determine file type
            ext = os.path.splitext(filepath)[1].lower()
            
            if ext == '.txt':
                # Export as plain text
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(self.text)
                messagebox.showinfo("Export Successful", f"Transcription saved to:\n{filepath}")
            
            elif ext == '.docx':
                # Export as Word document
                try:
                    from docx import Document
                    
                    doc = Document()
                    doc.add_heading(f'Transcription: {self.filename}', 0)
                    doc.add_paragraph(self.text)
                    doc.save(filepath)
                    
                    messagebox.showinfo("Export Successful", f"Transcription saved to:\n{filepath}")
                except ImportError:
                    messagebox.showerror(
                        "Missing Dependency",
                        "python-docx is required for .docx export.\n\n"
                        "Install it with: pip install python-docx\n\n"
                        "Falling back to .txt export..."
                    )
                    # Fallback to txt
                    txt_path = filepath.replace('.docx', '.txt')
                    with open(txt_path, 'w', encoding='utf-8') as f:
                        f.write(self.text)
                    messagebox.showinfo("Export Successful", f"Transcription saved to:\n{txt_path}")
            
            else:
                messagebox.showerror("Error", "Unsupported file format")
        
        except Exception as e:
            messagebox.showerror("Export Failed", f"Failed to export transcription:\n{str(e)}")
